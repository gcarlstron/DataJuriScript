import markdown
from bs4 import BeautifulSoup
from docx.shared import RGBColor
from docx import Document


class TestGerarDocx:

    def adicionar_tabela_docx(self, doc, dados_tabela: dict[str, str]):
        for para in doc.paragraphs:
            if '[[ tabela ]]' in para.text:

                para.text = para.text.replace('[[ tabela ]]', '')
                tabela = doc.add_table(rows=len(dados_tabela) + 1, cols=2)
                tabela.style = 'Table Grid'

                cabecalhos = ['Campo', 'Valor']
                for i, cabecalho in enumerate(cabecalhos):
                    tabela.cell(0, i).text = cabecalho

                for i, dicionario in enumerate(dados_tabela):
                    for chave, valor in dicionario.items():
                        tabela.cell(i, 0).text = chave
                        tabela.cell(i, 1).text = valor

                tabela.add_row()
                break

    def convert_markdown_to_docx(self, markdown_file, output_file, dados_tabela, styles=None):

        with open(markdown_file, 'r', encoding='utf-8') as file:
            documento = file.read()

        doc = Document()
        html = markdown.markdown(documento)
        soup = BeautifulSoup(html, 'html.parser')

        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol']):
            if element.name.startswith('h'):
                p = doc.add_paragraph(element.text)
                p.style = f'Heading {element.name[1]}'
            elif element.name == 'p':
                if '[[ tabela ]]' in element.text:
                    doc.add_paragraph(element.text)
                    self.adicionar_tabela_docx(doc, dados_tabela)
                else:
                    for child in element.children:
                        p = doc.add_paragraph()
                        if isinstance(child, str):
                            run = p.add_run(child)
                        else:
                            run = p.add_run(child.text)
                            if child.name == 'strong':
                                run.bold = True
                            elif child.name == 'em':
                                run.italic = True
                            elif child.name == 'a':
                                run.font.color.rgb = RGBColor(0, 0, 255)
                                run.underline = True
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    estilo = 'List Bullet' if element.name == 'ul' else 'List Number'
                    doc.add_paragraph(li.text, style=estilo)

        doc.save(output_file)


if __name__ == "__main__":
    mdfile(test.md)
    doc = Document()
    doc.add_paragraph("[[ tabela ]]")
    dados = [{'data_inicio': '15/06/1997', 'data_final': '01/08/2000', 'empresa': '<a  href="#/lista/PessoaJuridica/12173"  >Viação Aérea São Paulo S/A - VASP</a>', 'funcao': 'Comissária de Bordo', 'agentes_nocivos': 'Enquadramento profissional (cód. 2.4.1 do Decreto 53.831/64),  Ruído (cód. 2.0.1 Dec. 3.048/99),  Radiação ionizante (cód. 2.0.3 Dec. 3.048/99 e Grupo 1 LINACH),  Risco à integridade física (art. 57 Lei 8.213/91 c/c Súmula 198/TFR),  Vibrações (cód. 2.0.2 Dec. 3.048/99),  Penosidade (art. 57 Lei 8.213/91 c/c Súmula 198/TFR), ', 'provas': ''}]
    test = TestGerarDocx()
    test.adicionar_tabela_docx(doc, dados)
    test.convert_markdown_to_docx('test.md', 'test.docx', dados)
