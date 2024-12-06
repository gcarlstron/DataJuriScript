import json
import markdown
from io import BytesIO
from bs4 import BeautifulSoup
from docx.shared import RGBColor
from dotenv import load_dotenv
from docx import Document
from src.DataJuriAuthenticate import DataJuriAuthenticate
from src.DataJuriClient import DataJuriClient
from src.PreencherTemplate import arquivo_md



def main():
    # Configurações
    HOST = "api.datajuri.com.br"
    PROCESSO_ID = input("Digite o numero do processo: ")
    
    try:
        # Inicializa o cliente
        authenticate = DataJuriAuthenticate(HOST)
        client = DataJuriClient(HOST, authenticate.get_token())
        
        # Busca e preenche o template
        dados = client.preencher_template(PROCESSO_ID)
        documento, dados_tabela = arquivo_md('../template/retirement-request-template.md', dados)
        convert_markdown_to_docx(f'{documento}', documento.replace(".md", ".docx"), dados_tabela)
        
        # Salva o resultado em um arquivo
        with open('../template/dados_template.json', 'w', encoding='utf-8') as f:
            json.dump(dados, f, ensure_ascii=False, indent=2)
            
        print("Template preenchido com sucesso!")
        print("\nDados obtidos:")
        print(json.dumps(dados, ensure_ascii=False, indent=2))
        
    except Exception as e:
        print(f"Erro: {e}")


def adicionar_tabela_docx(doc, dados_tabela: dict[str, str]):
    for para in doc.paragraphs:
        if '[[ tabela ]]' in para.text:

            para.text = para.text.replace('[[ tabela ]]', '')
            tabela = doc.add_table(rows=len(dados_tabela) + 1, cols=2)
            tabela.style = 'Table Grid'

            cabecalhos = ['Campo', 'Valor']
            for i, cabecalho in enumerate(cabecalhos):
                tabela.cell(0, i).text = cabecalho

            for i, dicionario in enumerate(dados_tabela):
                for chave, valor in dicionario.items():
                    tabela.cell(i, 0).text = chave
                    tabela.cell(i, 1).text = valor

            tabela.add_row()
            break


def convert_markdown_to_docx(markdown_file, output_file, dados_tabela, styles=None):

    with open(markdown_file, 'r', encoding='utf-8') as file:
        documento = file.read()

    doc = Document()
    html = markdown.markdown(documento)
    soup = BeautifulSoup(html, 'html.parser')

    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol']):
        if element.name.startswith('h'):
            p = doc.add_paragraph(element.text)
            p.style = f'Heading {element.name[1]}'
        elif element.name == 'p':
            if '[[ tabela ]]' in element.text:
                # Adicionar a tabela
                adicionar_tabela_docx(doc, dados_tabela)
            else:
                p = doc.add_paragraph()
                for child in element.children:
                    if isinstance(child, str):
                        run = p.add_run(child)
                    else:
                        run = p.add_run(child.text)
                        if child.name == 'strong':
                            run.bold = True
                        elif child.name == 'em':
                            run.italic = True
                        elif child.name == 'a':
                            run.font.color.rgb = RGBColor(0, 0, 255)
                            run.underline = True
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li'):
                estilo = 'List Bullet' if element.name == 'ul' else 'List Number'
                doc.add_paragraph(li.text, style=estilo)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()


if __name__ == "__main__":
    load_dotenv()
    main()
