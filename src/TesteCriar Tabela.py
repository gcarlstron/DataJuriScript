from docx import Document


def adicionar_tabela_docx(doc, dados_tabela: [dict[str, str]]):
    for para in doc.paragraphs:
        if '[[ tabela ]]' in para.text:

            para.text = para.text.replace('[[ tabela ]]', '')
            tabela = doc.add_table(rows=len(dados_tabela) + 1, cols=2)
            tabela.style = 'Table Grid'

            cabecalhos = ['Campo', 'Valor']
            for i, cabecalho in enumerate(cabecalhos):
                tabela.cell(0, i).text = cabecalho

            for i, dicionario in enumerate(dados_tabela):
                for chave, valor in dicionario.items():
                    tabela.cell(i, 0).text = chave
                    tabela.cell(i, 1).text = valor

            tabela.add_row()
            break


if __name__ == "__main__":
    doc = Document()
    doc.add_paragraph("[[ tabela ]]")
    dados = [{'data_inicio': '15/06/1997', 'data_final': '01/08/2000', 'empresa': '<a  href="#/lista/PessoaJuridica/12173"  >Viação Aérea São Paulo S/A - VASP</a>', 'funcao': 'Comissária de Bordo', 'agentes_nocivos': 'Enquadramento profissional (cód. 2.4.1 do Decreto 53.831/64),  Ruído (cód. 2.0.1 Dec. 3.048/99),  Radiação ionizante (cód. 2.0.3 Dec. 3.048/99 e Grupo 1 LINACH),  Risco à integridade física (art. 57 Lei 8.213/91 c/c Súmula 198/TFR),  Vibrações (cód. 2.0.2 Dec. 3.048/99),  Penosidade (art. 57 Lei 8.213/91 c/c Súmula 198/TFR), ', 'provas': ''}]
    adicionar_tabela_docx(doc, dados)
