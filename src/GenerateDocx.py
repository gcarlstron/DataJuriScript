from typing import List, Dict, Any, Optional
import markdown
from bs4 import BeautifulSoup
from docx.shared import RGBColor, Pt, Cm
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
import re
from dataclasses import dataclass


@dataclass
class DocumentStyles:
    heading_styles: Dict[str, str] = None
    link_color: RGBColor = RGBColor(0, 0, 255)
    table_style: str = 'Table Grid'
    default_font: str = 'Cambria'
    heading_color: RGBColor = RGBColor(0, 0, 0)
    default_font_size: Pt = Pt(11)
    italic_left_indent: float = Cm(0.5)
    header_image: str = "../img/logo.jpg"
    image_width: float = Cm(5)


class GenerateDocx:
    def __init__(self, styles: Optional[DocumentStyles] = None):
        self.styles = styles or DocumentStyles()

    def _create_element(self, name: str) -> OxmlElement:
        return OxmlElement(name, xmlns=ns.w)

    def _create_header_with_image(self, doc: Document) -> None:
        if not self.styles.header_image:
            return

        section = doc.sections[0]

        header = section.header

        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = header_para.add_run()
        run.add_picture(self.styles.header_image, width=self.styles.image_width)

    def _clean_html_tags(self, text: str) -> str:
        """Remove HTML tags from text while preserving content."""
        return re.sub('<[^>]*>', '', text)

    def _apply_default_font_style(self, run) -> None:
        """Apply default font styling to a run."""
        run.font.name = self.styles.default_font
        run.font.size = self.styles.default_font_size

    def add_table_docx(self, doc: Document, table_data: Dict[str, str]) -> None:
        """Add a table to the document at the [[ tabela ]] marker."""
        for para in doc.paragraphs:
            if '[[ tabela ]]' not in para.text:
                continue

            para.text = para.text.replace('[[ tabela ]]', '')
            table = doc.add_table(rows=1, cols=2)
            table.style = self.styles.table_style

            headers = ['Campo', 'Valor']
            for i, header in enumerate(headers):
                cell = table.cell(0, i)
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        self._apply_default_font_style(run)

            for i, (key, value) in enumerate(table_data.items(), start=1):
                row = table.add_row()
                row.cells[0].text = key
                row.cells[1].text = self._clean_html_tags(value)
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            self._apply_default_font_style(run)
            break

    def _process_paragraph_element(self, doc: Document, element: BeautifulSoup) -> None:
        paragraph = doc.add_paragraph()
        has_italic = False

        for child in element.children:
            if isinstance(child, str):
                run = paragraph.add_run(child)
                self._apply_default_font_style(run)
                continue

            run = paragraph.add_run(child.text)
            self._apply_default_font_style(run)

            if child.name == 'strong':
                run.bold = True
            elif child.name == 'em':
                run.italic = True
                has_italic = True
            elif child.name == 'a':
                run.font.color.rgb = self.styles.link_color
                run.underline = True

        if has_italic:
            paragraph.paragraph_format.left_indent = self.styles.italic_left_indent

    def _format_heading(self, paragraph, text: str) -> None:
        run = paragraph.add_run(text.upper())
        run.font.name = self.styles.default_font
        run.font.size = self.styles.default_font_size
        run.font.color.rgb = self.styles.heading_color
        run.underline = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def convert_markdown_to_docx(
            self,
            markdown_file: str,
            output_file: str,
            table_data: List[Dict[str, str]]
    ) -> None:
        with open(markdown_file, 'r', encoding='utf-8') as file:
            content = file.read()

        doc = Document()

        self._create_header_with_image(doc)

        html = markdown.markdown(content)
        soup = BeautifulSoup(html, 'html.parser')
        table_index = 0

        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol']):
            if element.name.startswith('h'):
                p = doc.add_paragraph()
                p.style = f'Heading {element.name[1]}'
                self._format_heading(p, element.text)
            elif element.name == 'p':
                if '[[ tabela ]]' in element.text:
                    p = doc.add_paragraph(element.text)
                    for run in p.runs:
                        self._apply_default_font_style(run)
                    self.add_table_docx(doc, table_data[table_index])
                    table_index += 1
                else:
                    self._process_paragraph_element(doc, element)
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    style = 'List Bullet' if element.name == 'ul' else 'List Number'
                    p = doc.add_paragraph(li.text, style=style)
                    for run in p.runs:
                        self._apply_default_font_style(run)

        doc.save(output_file)
