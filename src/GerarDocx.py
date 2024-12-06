import markdown
from bs4 import BeautifulSoup
from docx.shared import RGBColor
from docx import Document


class GerarDocx:

    def adicionar_tabela_docx(self, doc, dados_tabela: dict[str, str]):
        for para in doc.paragraphs:
            if '[[ tabela ]]' in para.text:

                para.text = para.text.replace('[[ tabela ]]', '')
                tabela = doc.add_table(rows=len(dados_tabela) + 1, cols=2)
                tabela.style = 'Table Grid'

                cabecalhos = ['Campo', 'Valor']
                for i, cabecalho in enumerate(cabecalhos):
                    tabela.cell(0, i).text = cabecalho

                for i, dicionario in enumerate(dados_tabela):
                    for chave, valor in dicionario.items():
                        tabela.cell(i, 0).text = chave
                        tabela.cell(i, 1).text = valor

                tabela.add_row()
                break

    def convert_markdown_to_docx(self, markdown_file, output_file, dados_tabela, styles=None):

        with open(markdown_file, 'r', encoding='utf-8') as file:
            documento = file.read()

        doc = Document()
        html = markdown.markdown(documento)
        soup = BeautifulSoup(html, 'html.parser')

        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol']):
            if element.name.startswith('h'):
                p = doc.add_paragraph(element.text)
                p.style = f'Heading {element.name[1]}'
            elif element.name == 'p':
                if '[[ tabela ]]' in element.text:
                    doc.add_paragraph(element.text)
                    self.adicionar_tabela_docx(doc, dados_tabela)
                else:
                    for child in element.children:
                        p = doc.add_paragraph()
                        if isinstance(child, str):
                            run = p.add_run(child)
                        else:
                            run = p.add_run(child.text)
                            if child.name == 'strong':
                                run.bold = True
                            elif child.name == 'em':
                                run.italic = True
                            elif child.name == 'a':
                                run.font.color.rgb = RGBColor(0, 0, 255)
                                run.underline = True
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    estilo = 'List Bullet' if element.name == 'ul' else 'List Number'
                    doc.add_paragraph(li.text, style=estilo)

        doc.save(output_file)
