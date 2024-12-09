import os
import logging
from jinja2 import Template
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import markdown
from bs4 import BeautifulSoup
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table as RLTable, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib import colors
from reportlab.lib.units import inch

# Configuração do logging
logging.basicConfig(
    level=logging.DEBUG, 
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler() 
    ]
)

class PreenchimentoAutomaticoDocumento:
    def __init__(self, diretorio_templates):
        self.diretorio_templates = os.path.abspath(diretorio_templates)
        self.templates = {}
        self.carregar_templates()

    def carregar_templates(self):
        if not os.path.exists(self.diretorio_templates):
            logging.error(f"O diretório de templates não existe: {self.diretorio_templates}")
        
        for arquivo in os.listdir(self.diretorio_templates):
            if arquivo.endswith('.md') or arquivo.endswith('.txt'):
                caminho = os.path.join(self.diretorio_templates, arquivo)
                try:
                    with open(caminho, 'r', encoding='utf-8') as f:
                        conteudo = f.read()
                        nome_template = os.path.splitext(arquivo)[0]
                        self.templates[nome_template] = Template(conteudo)
                except Exception as e:
                    logging.error(f"Erro ao carregar o template {arquivo}: {str(e)}")
        
        if not self.templates:
            raise FileNotFoundError(f"Nenhum arquivo de template (.md ou .txt) encontrado em: {self.diretorio_templates}")

    def preencher_documento(self, nome_template, dados):
        if nome_template not in self.templates:
            raise ValueError(f"Template '{nome_template}' não encontrado. Templates disponíveis: {', '.join(self.templates.keys())}")
        
        template = self.templates[nome_template]
        dados['data_atual'] = datetime.now().strftime("%d de %B de %Y")
        
        # Formatar valores monetários
        for chave, valor in dados.items():
            if isinstance(valor, (int, float)) and 'valor' in chave.lower():
                dados[chave] = f"R$ {valor:.2f}"
        
        try:
            documento_preenchido = template.render(dados)
            if '[[ tabela ]]' in documento_preenchido:
                logging.debug("Marcador '[[ tabela ]]' está presente no documento preenchido.")
            else:
                logging.warning("Marcador '[[ tabela ]]' NÃO está presente no documento preenchido.")
            return documento_preenchido
        except Exception as e:
            logging.error(f"Erro ao renderizar o template '{nome_template}': {str(e)}")
            raise

    def adicionar_tabela_docx(self, doc, dados_tabela):
        marcador_encontrado = False
        for para in doc.paragraphs:
            if '[[ tabela ]]' in para.text:
                marcador_encontrado = True

                para.text = para.text.replace('[[ tabela ]]', '')
                tabela = doc.add_table(rows=len(dados_tabela) + 2, cols=5)
                tabela.style = 'Table Grid'
                
                cabecalhos = ['Item', 'Especificação', 'Und.', 'Preço Unitário R$', 'Preço Global R$']
                for i, cabecalho in enumerate(cabecalhos):
                    tabela.cell(0, i).text = cabecalho

                total_preco_global = 0
    
                for i, linha in enumerate(dados_tabela, 1):  # Começar de 1 para pular os cabeçalhos
                    tabela.cell(i, 0).text = str(linha['item'])
                    tabela.cell(i, 1).text = linha['especificacao']
                    tabela.cell(i, 2).text = str(linha['unidade'])
                    tabela.cell(i, 3).text = ""
                    tabela.cell(i, 4).text = f"R$ {linha['preco_global']:.2f}"
    
                    total_preco_global += linha['preco_global']
                
                total_linha = tabela.add_row()
                total_linha.cells[0].text = 'TOTAL MÁXIMO ESTIMADO'
                total_linha.cells[4].text = f"Total R$ {total_preco_global:.2f}"
                break 
    
        if not marcador_encontrado:
            logging.warning("Marcador '[[ tabela ]]' não encontrado no documento DOCX")

    def adicionar_tabela_pdf(self, flowables, dados_tabela):
        cabecalhos = ['Item', 'Especificação', 'Und.', 'Preço Unitário R$', 'Preço Global R$']
        dados = [cabecalhos] 
    
        total_preco_global = 0
    
        for linha in dados_tabela:
            dados.append([
                str(linha['item']),
                linha['especificacao'],
                str(linha['unidade']),
                "", 
                f"R$ {linha['preco_global']:.2f}"
            ])
    
        dados.append([
            'TOTAL MÁXIMO ESTIMADO', '', '', '', f"Total R$ {total_preco_global:.2f}"
        ])
    
        tabela = RLTable(dados, colWidths=[0.5 * inch, 2.5 * inch, 0.5 * inch, 1.5 * inch, 1.5 * inch])
        
        estilo_tabela = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey), 
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black), 
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),  
            ('GRID', (0, 0), (-1, -1), 1, colors.black), 
            ('BACKGROUND', (0, -1), (-1, -1), colors.lightgrey)  
        ])
        
        tabela.setStyle(estilo_tabela)
        flowables.append(tabela)
        flowables.append(Spacer(1, 12))

    def gerar_documento_txt(self, conteudo):
        conteudo = conteudo.replace('[[ tabela ]]', '[Tabela inserida aqui]')
        return conteudo.encode('utf-8')

    def gerar_documento_docx(self, conteudo, dados_tabela):
        doc = Document()
        html = markdown.markdown(conteudo)
        soup = BeautifulSoup(html, 'html.parser')
    
        # Iterar sobre os elementos e adicionar ao documento
        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol']):
            if element.name.startswith('h'):
                p = doc.add_paragraph(element.text)
                p.style = f'Heading {element.name[1]}'
            elif element.name == 'p':
                if '[[ tabela ]]' in element.text:
                    # Adicionar a tabela
                    self.adicionar_tabela_docx(doc, dados_tabela)
                else:
                    p = doc.add_paragraph()
                    for child in element.children:
                        if isinstance(child, str):
                            run = p.add_run(child)
                        else:
                            run = p.add_run(child.text)
                            if child.name == 'strong':
                                run.bold = True
                            elif child.name == 'em':
                                run.italic = True
                            elif child.name == 'a':
                                run.font.color.rgb = RGBColor(0, 0, 255)
                                run.underline = True
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    estilo = 'List Bullet' if element.name == 'ul' else 'List Number'
                    doc.add_paragraph(li.text, style=estilo)
    
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue()

    def gerar_documento_pdf(self, conteudo, dados_tabela):
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='Justify', alignment=TA_JUSTIFY))
        styles.add(ParagraphStyle(name='Centered', alignment=TA_CENTER))

        flowables = []
        html = markdown.markdown(conteudo, extensions=['extra'])
        soup = BeautifulSoup(html, 'html.parser')

        for element in soup.find_all(['div', 'p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'table']):
            centralizado = False
            parent_div = element.find_parent('div', align='center')
            if parent_div:
                centralizado = True
                
            if element.name == 'div' and element.get('align') == 'center':
                continue

            if element.name.startswith('h'):
                style = styles.get(f'Heading{element.name[1]}', styles['Heading1'])
                if centralizado:
                    style.alignment = TA_CENTER
                flowables.append(Paragraph(element.text, style))

            elif element.name == 'p':
                if '[[ tabela ]]' in element.text:
                    element_text = element.text.replace('[[ tabela ]]', '').strip()
                    if element_text:
                        if centralizado:
                            flowables.append(Paragraph(element_text, styles['Centered']))
                        else:
                            flowables.append(Paragraph(element_text, styles['Justify']))
                            
                    self.adicionar_tabela_pdf(flowables, dados_tabela)
                    flowables.append(Spacer(1, 12))
                else:
                    text = ''.join(str(child) for child in element.children)
                    text = text.replace('<strong>', '<b>').replace('</strong>', '</b>')
                    text = text.replace('<em>', '<i>').replace('</em>', '</i>')
                    text = text.replace('<a href="', '<link href="')  # Para links
                    if centralizado:
                        style = styles['Centered']
                    else:
                        style = styles['Justify']
                    flowables.append(Paragraph(text, style))

            elif element.name in ['ul', 'ol']:
                if element.name == 'ul':
                    for li in element.find_all('li'):
                        texto = f"• {li.text}"
                        if centralizado:
                            style = styles['Centered']
                        else:
                            style = styles['Normal']
                        flowables.append(Paragraph(texto, style))
                else:  # 'ol'
                    for i, li in enumerate(element.find_all('li'), 1):
                        texto = f"{i}. {li.text}"
                        if centralizado:
                            style = styles['Centered']
                        else:
                            style = styles['Normal']
                        flowables.append(Paragraph(texto, style))

            flowables.append(Spacer(1, 12))

        doc.build(flowables)
        buffer.seek(0)
        return buffer.getvalue()



    def processar_documento(self, nome_template, dados, formato, dados_tabela=None):
        documento_preenchido = self.preencher_documento(nome_template, dados)
        
        if formato == 'txt':
            return self.gerar_documento_txt(documento_preenchido)
        elif formato == 'docx':
            if dados_tabela:
                return self.gerar_documento_docx(documento_preenchido, dados_tabela)
            else:
                return self.gerar_documento_docx(documento_preenchido, [])
        elif formato == 'pdf':
            if dados_tabela:
                return self.gerar_documento_pdf(documento_preenchido, dados_tabela)
            else:
                logging.warning("Nenhum dado de tabela fornecido para PDF")
                return self.gerar_documento_pdf(documento_preenchido, [])
        else:
            logging.error(f"Formato de arquivo não suportado: {formato}")
            raise ValueError("Formato de arquivo não suportado. Use 'txt', 'docx' ou 'pdf'")

#testando a funcao
if __name__ == "__main__":
    try:
        preenchedor = PreenchimentoAutomaticoDocumento("./templates")
        
        dados = {
            "processo_administrativo": "123456789",
            "pregao_eletronico": "001",
            "ano_dispensa": "2024",
            "nome_orgao": "Nome do Órgão",
            "numero_dispensa": "001",
            "ano_dispensa": "2024",
            "nome_empresa": "LCT LICITACOES EPP LIMITADA",
            "endereco": "Loteamento Conviver Parnaíba Residence nº 9156, quadra 6, casa 19, bairro Joao XXIII",
            "cep": "64205-030",
            "municipio": "Parnaíba",
            "uf": "PI",
            "cnpj": "55.971.563/0001-19",
            "inscricao_estadual_municipal": "197664628",
            "prazo_validade_proposta": "60",
            "prazo_validade_proposta_extenso": "sessenta",
            "razao_social": "LCT LICITACOES EPP LIMITADA",
            "endereco_completo": "Loteamento Conviver Parnaíba Residence nº 9156, quadra 6, casa 19, bairro Joao XXIII, CEP: 64205-030, Parnaíba-PI",
            "telefone": "86-981281928",
            "email": "LCTLICITACOES@GMAIL.COM",
            "dados_bancarios": "Banco: XXX, Agência: YYYY, Conta: ZZZZ",
            "nome_representante": "Luan Yuuki Brito Asano",
            "endereco_representante": "BR 343 Conviver Parnaíba Residence",
            "cep_representante": "64205-030",
            "cidade_representante": "Parnaíba",
            "uf_representante": "PI",
            "cpf_representante": "058.188.063-32",
            "rg_representante": "2008283174-7",
            "cargo_representante": "Sócio Presidente",
            "orgao_expedidor_rg": "SSP-CE",
            "naturalidade_representante": "Fortaleza - CE",
            "nacionalidade_representante": "brasileiro",
            "dia_local": "24",
            "mes_local": "setembro",
            "ano_local": "2024",
            "valor_item": "1.000,00"  
        }
        
        dados_tabela = [
            {"item": 1, "especificacao": "(produto do orçamento do fornecedor)", "unidade": 1, "preco_unitario": 100.00, "preco_global": 500.00},
            {"item": 2, "especificacao": "(produto do orçamento do fornecedor)", "unidade": 1, "preco_unitario": 200.00, "preco_global": 600.00},
            {"item": 3, "especificacao": "(produto do orçamento do fornecedor)", "unidade": 1, "preco_unitario": 300.00, "preco_global": 700.00}
        ]
        
        for formato in ['txt', 'docx', 'pdf']:
            try:
                logging.info(f"Iniciando a geração do documento no formato {formato.upper()}")
                content = preenchedor.processar_documento(
                    nome_template="proposta_pregao", 
                    dados=dados, 
                    formato=formato, 
                    dados_tabela=dados_tabela 
                )
                logging.info(f"{formato.upper()} gerado com sucesso!")
                with open(f"documento.{formato}", "wb") as f:
                    f.write(content)
                logging.debug(f"Arquivo {f.name} salvo com sucesso")
            except Exception as e:
                logging.error(f"Erro ao gerar {formato.upper()}: {str(e)}")
    
    except Exception as e:
        logging.critical(f"Erro ao inicializar o preenchedor: {str(e)}")
        logging.debug(f"Diretório atual: {os.getcwd()}")
        logging.debug("Conteúdo do diretório 'templates':")
        try:
            for file in os.listdir("./templates"):
                logging.debug(f"  - {file}")
        except FileNotFoundError:
            logging.error("  O diretório 'templates' não foi encontrado.")
